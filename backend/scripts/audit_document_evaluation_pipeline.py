"""End-to-end audit of the Professional Document Evaluation pipeline.

Run: python -m scripts.audit_document_evaluation_pipeline
"""

from __future__ import annotations

import io
import sys
import tempfile
import traceback
from datetime import datetime
from pathlib import Path
from unittest.mock import MagicMock, patch

# Ensure backend root is importable
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.agents.professional_document_evaluation_agent import professional_document_evaluation_agent
from app.agents.report_generation_agent import report_generation_agent
from app.agents.state import AnalysisGraphState
from app.agents.workflow import create_analysis_workflow
from app.chains.extraction_parser import (
    ApplicantProfile,
    CVFields,
    ExtractedDocument,
    LanguageTestFields,
    PassportFields,
    TranscriptFields,
)
from app.chains.output_parser import ApplicationAnalysisResult
from app.models import Application, Document
from app.services.document_content_service import DocumentContentError, extract_text_from_bytes
from app.services.document_evaluation.evaluators.base import find_extracted_document
from app.services.pipeline_context import AnalysisPipelineContext
from app.services.readiness_report_builder import (
    _build_document_assessment,
    _build_document_assessment_from_evaluations,
    _build_document_assessment_legacy,
    build_readiness_report,
)

PASS = "PASS"
FAIL = "FAIL"
WARN = "WARN"

results: list[dict] = []


def record(step: str, status: str, detail: str, *, file: str = "", func: str = "") -> None:
    results.append({"step": step, "status": status, "detail": detail, "file": file, "func": func})
    icon = {"PASS": "+", "FAIL": "X", "WARN": "!"}.get(status, "?")
    loc = f" [{file}::{func}]" if file else ""
    print(f"  [{icon}] {step}: {detail}{loc}")


def _application() -> Application:
    return Application(
        id="audit-app",
        user_id="user-1",
        application_type="scholarship",
        status="analyzing",
        title="Audit Scholarship",
        country="UK",
        source_url=None,
        readiness_score=None,
        created_at=datetime(2026, 1, 1),
    )


def _doc(doc_id: str, doc_type: str, file_name: str) -> Document:
    return Document(
        id=doc_id,
        application_id="audit-app",
        file_name=file_name,
        document_type=doc_type,
        storage_path=f"/docs/{file_name}",
        uploaded_at=datetime(2026, 1, 15),
        file_size=1024,
        mime_type="application/pdf",
    )


def audit_workflow_graph() -> None:
    print("\n=== Step 1: LangGraph workflow includes document evaluation agent ===")
    workflow = create_analysis_workflow()
    graph = workflow.get_graph()
    nodes = set(graph.nodes.keys())
    edges = {(e[0], e[1]) for e in graph.edges}

    has_prof_eval = "professional_document_evaluation" in nodes
    has_legacy_agent = "document_evaluation" in nodes

    if has_prof_eval:
        record(
            "workflow_node",
            PASS,
            f"Node 'professional_document_evaluation' present; nodes={sorted(nodes)}",
            file="app/agents/workflow.py",
            func="create_analysis_workflow",
        )
    else:
        record(
            "workflow_node",
            FAIL,
            f"Missing professional_document_evaluation node; nodes={sorted(nodes)}",
            file="app/agents/workflow.py",
            func="create_analysis_workflow",
        )

    expected_edge = ("requirement_matching", "professional_document_evaluation")
    if expected_edge in edges:
        record(
            "workflow_edge",
            PASS,
            "Edge requirement_matching -> professional_document_evaluation exists",
            file="app/agents/workflow.py",
            func="create_analysis_workflow",
        )
    else:
        record(
            "workflow_edge",
            FAIL,
            f"Missing edge {expected_edge}; edges={sorted(edges)}",
            file="app/agents/workflow.py",
            func="create_analysis_workflow",
        )

    if has_legacy_agent:
        record(
            "legacy_agent_node",
            WARN,
            "Legacy 'document_evaluation' node also present (duplicate?)",
            file="app/agents/workflow.py",
            func="create_analysis_workflow",
        )
    else:
        record(
            "document_evaluation_agent",
            WARN,
            "document_evaluation_agent is NOT wired into workflow; "
            "professional_document_evaluation_agent is used instead "
            "(identical implementation in app/agents/document_evaluation_agent.py)",
            file="app/agents/workflow.py",
            func="create_analysis_workflow",
        )


def audit_agent_populates_ctx() -> None:
    print("\n=== Step 2: ctx.document_evaluations populated after evaluation ===")
    ctx = AnalysisPipelineContext(application_id="audit-app", application=_application())
    ctx.documents = [_doc("doc-cv", "cv", "cv.pdf"), _doc("doc-passport", "passport", "passport.pdf")]
    ctx.applicant_profile = ApplicantProfile()
    ctx.requirements = []
    state: AnalysisGraphState = {"ctx": ctx, "llm": MagicMock()}

    with patch(
        "app.agents.professional_document_evaluation_agent.DocumentEvaluationService"
    ) as mock_cls:
        mock_svc = MagicMock()
        mock_evals = [
            MagicMock(document_id="doc-cv", document_type="cv"),
            MagicMock(document_id="doc-passport", document_type="passport"),
        ]
        mock_svc.evaluate_all.return_value = mock_evals
        mock_cls.return_value = mock_svc

        out = professional_document_evaluation_agent(state)

    evals = out["ctx"].document_evaluations
    if evals is not None and len(evals) == 2:
        record(
            "ctx.document_evaluations",
            PASS,
            f"Populated with {len(evals)} evaluation(s) after agent run",
            file="app/agents/professional_document_evaluation_agent.py",
            func="professional_document_evaluation_agent",
        )
    else:
        record(
            "ctx.document_evaluations",
            FAIL,
            f"Expected 2 evaluations, got {evals!r}",
            file="app/agents/professional_document_evaluation_agent.py",
            func="professional_document_evaluation_agent",
        )


def audit_report_generation_passes_evaluations() -> None:
    print("\n=== Step 3: document_evaluations passed into build_readiness_report() ===")
    captured: dict = {}

    def capture_build(*args, **kwargs):
        captured["document_evaluations"] = kwargs.get("document_evaluations")
        return MagicMock()

    ctx = AnalysisPipelineContext(application_id="audit-app", application=_application())
    ctx.retriever = MagicMock()
    ctx.application_information = "info"
    ctx.applicant_profile = ApplicantProfile()
    ctx.retrieval_query = "query"
    ctx.documents = [_doc("doc-cv", "cv", "cv.pdf")]
    mock_evals = [MagicMock(document_id="doc-cv", document_type="cv")]
    ctx.document_evaluations = mock_evals

    with (
        patch("app.agents.report_generation_agent.ApplicationAnalysisChain") as chain_cls,
        patch("app.agents.report_generation_agent.build_readiness_report", side_effect=capture_build),
    ):
        chain = MagicMock()
        chain.analyze.return_value = ApplicationAnalysisResult(readiness_score=70)
        chain_cls.return_value = chain
        report_generation_agent({"ctx": ctx, "llm": MagicMock()})

    passed = captured.get("document_evaluations")
    if passed is mock_evals:
        record(
            "build_readiness_report args",
            PASS,
            "ctx.document_evaluations forwarded to build_readiness_report()",
            file="app/agents/report_generation_agent.py",
            func="report_generation_agent",
        )
    else:
        record(
            "build_readiness_report args",
            FAIL,
            f"document_evaluations not forwarded correctly: {passed!r}",
            file="app/agents/report_generation_agent.py",
            func="report_generation_agent",
        )


def audit_build_document_assessment_routing() -> None:
    print("\n=== Step 4-5: _build_document_assessment routing ===")
    from app.schemas.document_evaluation import DocumentEvaluationResult

    documents = [_doc("doc-cv", "cv", "cv.pdf")]
    evaluations = [
        DocumentEvaluationResult(
            document_id="doc-cv",
            file_name="cv.pdf",
            document_type="cv",
            quality_score=85,
            completeness="Complete",
            quality_rating="Good",
            strengths=["Clear structure"],
            weaknesses=[],
            suggestions=[],
        )
    ]

    with patch(
        "app.services.readiness_report_builder._build_document_assessment_from_evaluations"
    ) as from_evals:
        with patch(
            "app.services.readiness_report_builder._build_document_assessment_legacy"
        ) as legacy:
            from_evals.return_value = [MagicMock()]
            legacy.return_value = [MagicMock()]
            _build_document_assessment(documents, None, evaluations)

    if from_evals.called and not legacy.called:
        record(
            "_build_document_assessment routing",
            PASS,
            "Non-empty document_evaluations -> _build_document_assessment_from_evaluations()",
            file="app/services/readiness_report_builder.py",
            func="_build_document_assessment",
        )
    else:
        record(
            "_build_document_assessment routing",
            FAIL,
            f"from_evaluations called={from_evals.called}, legacy called={legacy.called}",
            file="app/services/readiness_report_builder.py",
            func="_build_document_assessment",
        )

    with patch(
        "app.services.readiness_report_builder._build_document_assessment_from_evaluations"
    ) as from_evals2:
        with patch(
            "app.services.readiness_report_builder._build_document_assessment_legacy"
        ) as legacy2:
            legacy2.return_value = [MagicMock()]
            _build_document_assessment(documents, ApplicantProfile(), None)

    if legacy2.called and not from_evals2.called:
        record(
            "legacy fallback",
            PASS,
            "Empty document_evaluations -> _build_document_assessment_legacy()",
            file="app/services/readiness_report_builder.py",
            func="_build_document_assessment",
        )
    else:
        record(
            "legacy fallback",
            FAIL,
            f"Legacy path not used when evaluations=None",
            file="app/services/readiness_report_builder.py",
            func="_build_document_assessment",
        )


def audit_load_document_text() -> None:
    print("\n=== Step 6: load_document_text() for PDF/DOCX ===")

    # Minimal valid PDF (one blank page)
    pdf_bytes = (
        b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>stream\nBT /F1 12 Tf 100 700 Td (Hello CV) Tj ET\nendstream\nendobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n"
        b"0000000115 00000 n \n0000000261 00000 n \n0000000354 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n422\n%%EOF"
    )

    try:
        pdf_text = extract_text_from_bytes(pdf_bytes, file_name="test.pdf", mime_type="application/pdf")
        if pdf_text.strip():
            record(
                "PDF extraction",
                PASS,
                f"extract_text_from_bytes PDF ok, len={len(pdf_text.strip())}, snippet={pdf_text.strip()[:40]!r}",
                file="app/services/document_content_service.py",
                func="extract_text_from_bytes",
            )
        else:
            record(
                "PDF extraction",
                FAIL,
                "PDF extracted but text empty",
                file="app/services/document_content_service.py",
                func="extract_text_from_bytes",
            )
    except Exception as exc:
        record(
            "PDF extraction",
            FAIL,
            f"{type(exc).__name__}: {exc}",
            file="app/services/document_content_service.py",
            func="extract_text_from_bytes",
        )

    # DOCX via python-docx if available, else skip with note
    try:
        from docx import Document as DocxDocument

        buf = io.BytesIO()
        doc = DocxDocument()
        doc.add_paragraph("Sample motivation letter content for audit.")
        doc.save(buf)
        docx_bytes = buf.getvalue()
        docx_text = extract_text_from_bytes(
            docx_bytes,
            file_name="test.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        if "motivation letter" in docx_text.lower():
            record(
                "DOCX extraction",
                PASS,
                f"extract_text_from_bytes DOCX ok, snippet={docx_text.strip()[:50]!r}",
                file="app/services/document_content_service.py",
                func="extract_text_from_bytes",
            )
        else:
            record(
                "DOCX extraction",
                FAIL,
                f"DOCX text unexpected: {docx_text!r}",
                file="app/services/document_content_service.py",
                func="extract_text_from_bytes",
            )
    except ImportError:
        record(
            "DOCX extraction",
            WARN,
            "python-docx not installed; skipped DOCX runtime test",
            file="app/services/document_content_service.py",
            func="extract_text_from_bytes",
        )
    except Exception as exc:
        record(
            "DOCX extraction",
            FAIL,
            f"{type(exc).__name__}: {exc}",
            file="app/services/document_content_service.py",
            func="extract_text_from_bytes",
        )

    # load_document_text integration (mocked download)
    from app.services.document_content_service import load_document_text

    doc = _doc("doc-pdf", "cv", "cv.pdf")
    with patch(
        "app.services.document_content_service.download_document_bytes",
        return_value=pdf_bytes,
    ):
        try:
            loaded = load_document_text(doc)
            record(
                "load_document_text",
                PASS,
                f"load_document_text() returned {len(loaded)} chars via mocked download",
                file="app/services/document_content_service.py",
                func="load_document_text",
            )
        except DocumentContentError as exc:
            record(
                "load_document_text",
                FAIL,
                str(exc),
                file="app/services/document_content_service.py",
                func="load_document_text",
            )


def audit_find_extracted_document() -> None:
    print("\n=== Step 7: find_extracted_document() matching ===")
    profile = ApplicantProfile(
        documents=[
            ExtractedDocument(
                document_id="doc-cv",
                file_name="cv.pdf",
                document_type="cv",
                cv=CVFields(skills=["Python"]),
            ),
            ExtractedDocument(
                document_id="doc-other",
                file_name="passport.pdf",
                document_type="passport",
                passport=PassportFields(full_name="Jane Doe"),
            ),
        ]
    )

    by_id = find_extracted_document(profile, _doc("doc-cv", "cv", "cv.pdf"))
    if by_id and by_id.document_id == "doc-cv":
        record("match by document_id", PASS, "Matched doc-cv by document_id", file="app/services/document_evaluation/evaluators/base.py", func="find_extracted_document")
    else:
        record("match by document_id", FAIL, f"Expected doc-cv, got {by_id}", file="app/services/document_evaluation/evaluators/base.py", func="find_extracted_document")

    by_type = find_extracted_document(profile, _doc("doc-new", "passport", "passport2.pdf"))
    if by_type and by_type.document_type == "passport":
        record("match by document_type", PASS, "Matched passport by document_type fallback", file="app/services/document_evaluation/evaluators/base.py", func="find_extracted_document")
    else:
        record("match by document_type", FAIL, f"Type fallback failed: {by_type}", file="app/services/document_evaluation/evaluators/base.py", func="find_extracted_document")

    no_profile = find_extracted_document(None, _doc("doc-cv", "cv", "cv.pdf"))
    if no_profile is None:
        record("no profile", PASS, "Returns None when applicant_profile is None", file="app/services/document_evaluation/evaluators/base.py", func="find_extracted_document")
    else:
        record("no profile", FAIL, f"Expected None, got {no_profile}", file="app/services/document_evaluation/evaluators/base.py", func="find_extracted_document")


def audit_per_document_evaluation() -> None:
    print("\n=== Step 8: Per-document evaluation outcomes ===")
    from app.services.document_evaluation import DocumentEvaluationService

    service = DocumentEvaluationService(llm=None)
    sample_text = (
        "john@example.com\nExperience\nEducation\nSkills\nPython Machine Learning\n"
        "Improved deployment speed by 30%\n"
        "github.com/applicant\nlinkedin.com/in/applicant"
    )

    scenarios: list[tuple[str, Document, ExtractedDocument | None, str, str]] = [
        (
            "cv (success)",
            _doc("d1", "cv", "cv.pdf"),
            ExtractedDocument(
                document_id="d1",
                file_name="cv.pdf",
                document_type="cv",
                cv=CVFields(skills=["Python"], experience=["Intern"], leadership=["President"]),
            ),
            "success",
            "Rule-based CV evaluation with extracted fields and readable text",
        ),
        (
            "passport (success)",
            _doc("d2", "passport", "passport.pdf"),
            ExtractedDocument(
                document_id="d2",
                file_name="passport.pdf",
                document_type="passport",
                passport=PassportFields(
                    full_name="Jane Doe",
                    nationality="US",
                    passport_number="AB1234567",
                    expiry_date="2030-12-31",
                ),
            ),
            "success",
            "Passport fields extracted; expiry valid",
        ),
        (
            "academic_transcript (success)",
            _doc("d3", "academic_transcript", "transcript.pdf"),
            ExtractedDocument(
                document_id="d3",
                file_name="transcript.pdf",
                document_type="academic_transcript",
                transcript=TranscriptFields(
                    university="Cairo University",
                    degree="BSc",
                    major="CS",
                    gpa="3.8",
                    graduation_year="2024",
                ),
            ),
            "success",
            "Transcript fields populated",
        ),
        (
            "ielts_score (success)",
            _doc("d4", "ielts_score", "ielts.pdf"),
            ExtractedDocument(
                document_id="d4",
                file_name="ielts.pdf",
                document_type="ielts_score",
                language_test=LanguageTestFields(
                    test_type="IELTS",
                    overall_score="7.5",
                    reading="8.0",
                    listening="7.5",
                    writing="7.0",
                    speaking="7.5",
                ),
            ),
            "success",
            "IELTS component scores present",
        ),
        (
            "cv (skipped)",
            _doc("d5", "cv", "cv_scanned.pdf"),
            ExtractedDocument(
                document_id="d5",
                file_name="cv_scanned.pdf",
                document_type="cv",
                extraction_status="skipped",
                error_message="Image-only PDF",
            ),
            "skipped",
            "Extraction status skipped -> evaluator short-circuits",
        ),
        (
            "other/unreadable (failed-like: success with score 0)",
            _doc("d6", "other", "blank.pdf"),
            None,
            "success",
            "No text loadable -> GenericDocumentEvaluator returns score 0 (evaluation_status still 'success')",
        ),
    ]

    profile_docs = [s[2] for s in scenarios if s[2] is not None]
    profile = ApplicantProfile(documents=[d for d in profile_docs if d])

    for label, document, extracted, expected_status, reason in scenarios:
        if extracted and extracted.extraction_status == "skipped":
            text = None
        elif document.document_type == "other":
            text = None
        else:
            text = sample_text

        with patch("app.services.document_evaluation.evaluator.load_document_text") as mock_load:
            if text is not None:
                mock_load.return_value = text
            else:
                mock_load.side_effect = DocumentContentError("No extractable text")

            try:
                result = service.evaluate_document(
                    document,
                    applicant_profile=profile,
                    requirements=[],
                )
                actual = result.evaluation_status
                score = result.quality_score
                detail = (
                    f"{label}: status={actual}, score={score}, "
                    f"weaknesses={result.weaknesses[:2]}, reason={reason}"
                )
                if actual == expected_status:
                    record(f"eval:{label}", PASS, detail, file="app/services/document_evaluation/evaluator.py", func="evaluate_document")
                else:
                    record(
                        f"eval:{label}",
                        FAIL,
                        f"Expected status={expected_status}, got {actual}; {detail}",
                        file="app/services/document_evaluation/evaluators/base.py",
                        func="BaseDocumentEvaluator.evaluate",
                    )
            except Exception as exc:
                record(
                    f"eval:{label}",
                    FAIL,
                    f"{type(exc).__name__}: {exc}\n{traceback.format_exc()}",
                    file="app/services/document_evaluation/evaluator.py",
                    func="evaluate_document",
                )


def audit_end_to_end_report_integration() -> None:
    print("\n=== Step 9: End-to-end report uses professional evaluations ===")
    from app.schemas.document_evaluation import DocumentEvaluationResult

    evaluations = [
        DocumentEvaluationResult(
            document_id="doc-cv",
            file_name="cv.pdf",
            document_type="cv",
            quality_score=88,
            completeness="Complete",
            quality_rating="Good",
            strengths=["Strong projects"],
            weaknesses=["Add quantified achievements"],
            suggestions=["Add metrics to experience bullets"],
        )
    ]
    report = build_readiness_report(
        ApplicationAnalysisResult(readiness_score=72, missing_documents=["Passport"]),
        readiness_score=72,
        application=_application(),
        documents=[_doc("doc-cv", "cv", "cv.pdf")],
        requirements=[],
        document_evaluations=evaluations,
    )
    cv = next((a for a in report.document_assessment if a.document_type == "cv"), None)
    if cv and cv.quality_score == 88 and cv.quality == "88/100 (Good)":
        record(
            "report integration",
            PASS,
            "build_readiness_report produced professional CV assessment entry",
            file="app/services/readiness_report_builder.py",
            func="build_readiness_report",
        )
    else:
        record(
            "report integration",
            FAIL,
            f"CV assessment missing or wrong: {cv}",
            file="app/services/readiness_report_builder.py",
            func="build_readiness_report",
        )


def print_summary() -> int:
    print("\n" + "=" * 72)
    print("AUDIT SUMMARY")
    print("=" * 72)
    passed = sum(1 for r in results if r["status"] == PASS)
    failed = sum(1 for r in results if r["status"] == FAIL)
    warned = sum(1 for r in results if r["status"] == WARN)
    print(f"  PASS: {passed}  FAIL: {failed}  WARN: {warned}  TOTAL: {len(results)}")

    failures = [r for r in results if r["status"] == FAIL]
    if failures:
        print("\nFAILURES:")
        for f in failures:
            loc = f"{f['file']}::{f['func']}" if f["file"] else "unknown"
            print(f"  - {f['step']} @ {loc}: {f['detail']}")

    warnings = [r for r in results if r["status"] == WARN]
    if warnings:
        print("\nWARNINGS:")
        for w in warnings:
            print(f"  - {w['step']}: {w['detail']}")

    return 1 if failures else 0


def main() -> None:
    print("Professional Document Evaluation Pipeline Audit")
    print("=" * 72)
    audit_workflow_graph()
    audit_agent_populates_ctx()
    audit_report_generation_passes_evaluations()
    audit_build_document_assessment_routing()
    audit_load_document_text()
    audit_find_extracted_document()
    audit_per_document_evaluation()
    audit_end_to_end_report_integration()
    raise SystemExit(print_summary())


if __name__ == "__main__":
    main()
